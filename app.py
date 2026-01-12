import streamlit as st
import os
import re
import camelot
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import pandas as pd
import io
import base64
import subprocess
import streamlit.components.v1 as components

# --- 內置 Word 模板 (Base64) ---
# 採用列表形式存儲，避免單行字符串過長導致的解析問題
TEMPLATE_CHUNKS = [
    "UEsDBAoAAAAAAIdO4kAAAAAAAAAAAAAAAAAJAAAAZG9jUHJvcHMvUEsDBBQAAAAIAIdO4kAvZypzfgEA",
    "AJkCAAAQAAAAZG9jUHJvcHMvYXBwLnhtbJ1STU/jMBS8I+1/iHJPbKcfpOjVqAQ4LWylBnpElvPSWJvY",
    "lm0Q/ffr0N02e+X2ZkYez/uA28+hTz7QeWX0OmU5TRPU0jRKH9bpS/2YlWnig9CN6I3GdXpEn97yH1ew",
    "dcaiCwp9Ei20X6ddCPaGEC87HITPo6yj0ho3iBChOxDTtkrivZHvA+pACkqXBD8D6gabzJ4N05PjzUf4",
    "rmlj5JjPv9ZHGwNzqHGwvQjIn8c4fd6YMAA5s1CZwQp95E9KOuNNG4D8o2ArDuj5DMipgL1xjeesKCmQ",
    "Uw1VJ5yQIY6Rs9VyFR9fCPipdHzP5kBOVTR04uCE7TyP5ARBbYLoazUgZ2UZ850h7KTosYo98Vb0HoFc",
    "iPGD3/7F1uZ+7PCv/j85SbhXodtZIcdMq5JNs04k2FjbKylCvAu+3+6SX1+7e2NFHo8kL+Z0xt4e2cOs",
    "uL6rsmK5qrL5bNFkG7YoMrqoFnNaUlpUGyBTJ4jL36F8dyoceZzfFMZRnE+A/wFQSwMEFAAAAAgAh07i",
    "QCCPaUt1AQAAswIAABEAAABkb2NQcm9wcy9jb3JlLnhtbH2SQU7DMBBF90jcIfI+sZ2EqlhpkACxolIl",
    "gqjYWfZQDIkT2Ya2N+AIvQdnYsEtcNIkUIFY2vPnzf8eZ2ebqgxewVhV6xmiEUEBaFFLpVczdFtchVMU",
    "WMe15GWtYYa2YNFZfnyUiYaJ2sDC1A0Yp8AGnqQtE80MPTrXMIyteISK28grtC8+1Kbizh/NCjdcPPMV",
    "4JiQCa7Acckdxy0wbEYi6pFSjMjmxZQdQAoMJVSgncU0ovhb68BU9s+GrvJDWSm3bXym3u5PthT74qje",
    "WDUK1+t1tE46G94/xcv59U0XNVS6fSsBKM+k6MYxYYA7kIEHsP24oXKXXFwWVyiPSZyGZBLStCCEpVNG",
    "yH2GB1Xf3wL3rNrky9tgrvTTi25l4227kpJbN/fbe1Agz7f5x9vuc/ee4d+VQbwwSnt7vYkkjNOCEpac",
    "7k30fYNoDFX1I/5PNQkJDWlckJjFyWGqAZC3Pgy8vb+fTkpPM6XnSnw2+WfwFQSwMEFAAAAAgAh07i",
    "QLmkklPiAQAABgQAABMAAABkb2NQcm9wcy9jdXN0b20ueG1stZNLj5swFIX3lfofkPcE24QkRCEjwCTK",
    "dCDNoxklmwqBwzgDNgXnNVX/e53mMZrFbNrO0rpXx989597e3aHItR2taia4A1ADAo3yRKSMZw74Nh/o",
    "HaDVMuZpnAtOHXCkNbjrf/7U+1qJklaS0VpTErx2wJOUZdcw6uSJFnHdUGWuKmtRFbFUzyozxHrNEkpE",
    "si0olwaGsGUk21qKQi9vcuCs193Jv5VMRXKiqxfzY6lw+72L+FFbF5KlDvhJLJ8QC1o6DmxfRxB5um3a",
    "bR12IMQe9ge2G/wCWnlqxkDjcaFG96djpbWT3bzc17Lq73/kD/GjxRNsF6nvTabBMkuHnWyBvPmIZPue",
    "8drbM64M/0hj3mgqGkuanonWLKeSFbSPITbVPDruzBHumu0uhKs/HLeO/0bSvJJ8mY2V3ek2kd6W5emC",
    "Vm9swtDCOsINtVsN3IQm+hBfrCvNyF+8+d4y/ZbpoXbgN82m7bpuCxKEB4HKHLZdjL8j/CFArSuQsmdO",
    "izJXaanNp9Us3tEpTUR1ye6yTfR4L1bD/Hm0ESzaJFZEBsVyk+fLl+Cw2mQwHEbPEQlxiCdoRSYofIlY",
    "RDL24N+jBC+Os8cJG7PRISRuMyTLw5gEcMSh895wxulEzgfc/w1QSwMECgAAAAAAh07iQAAAAAAAAAAA",
    "AAAAAAUAAAB3b3JkL1BLAwQUAAAACACHTuJAqODxmxENAAAlbwAADwAAAHdvcmQvc3R5bGVzLnhtbN1d",
    "zW/cuBW/F+j/IMypPTj22ONPxFk4ttMEtV2343TPHInjUSKJU0mTiXMOUBQt0EOBLbDYQ9tDT9u99bbt",
    "X9Mku/9FHylKw5FEavQkboHm4uiD7/f0Pn585FDU48/ehoHzhsaJz6LTwfDRzsChkcs8P7o/Hby8e7Z1",
    "NHCSlEQeCVhETwcPNBl89uTHP3q8PEnSh4AmDgiIkpPQPR3M0nR+sr2duDMakuQRm9MILk5ZHJIUDuP7",
    "7ZDErxfzLZeFc5L6Ez/w04ft3Z2dg4EUw04Hizg6kSK2Qt+NWcKmKW9ywqZT36XyT94i3gQ3a3nB3EVI",
    "o1Qgbsc0AB1YlMz8eZJLC7HS4BFnuZA3pod4Ewb5fctNwJYs9uYxc2mSgE/CIFM+JH5UiBmOKoIKwz0C",
    "w21nj7/NRUHz4Y74n6LHcMeksTQ7b51DJkEFscbbmRev/ElM4szNEACK3vPkfJGkLLwgKSnkLZfLR8t5",
    "8siNpNqK14Z723Bp1WjghO7Ji/uIxWQSQHAuh6PBE4hMj7kXdEoWQZrww/g2lofySPx5xqI0cZYnJHF9",
    "/3RwFvsEPLM8mZ1FiXpMSZKeJT65g6gGjNAHuEt5jt/vJnnj7SePtwVa/ldBnRc6ZHeVVIRQhMAcZxkF",
    "MtkiSk8HuweQjvA4dPrLZyKLTgf5iZfRzPfo5zMavUyoB5krbxzT0H/uex7l2SzPvXxxG/sshlw7HRwf",
    "y5NXzH1NvXEKwFwqN0mQeJdvXTrnWQGwv8kxhZxFCVAosvBXksWJRIEXJyLCjXbDtQ8G8Og2UWaUcN5y",
    "hpsAlTXPFM1F7HYXsdddxKi7iP3uIg66izjsLuKou4jjWhGdAtuPPPpWE3A9CK4Pwx4E1wdnD4LrQ7YH",
    "wfWB3IPg+vDuQXB90PcguD4VehBsIUFS5tpIDy7WQnJwsRZSg4u1kBhcrIW04GItJAUXayEluFgLCcHF",
    "WkiHrBByXkC3EaX9d0dTxtKIpdRJ6VsL4kkEwsWgSQ/QY9XIizEa11qpRxRuMw1KD5QqC8pNHqK+InWJ",
    "qMdrBXTSL+VjJodNnal/v4hhWF9XnXdCoNEbGsC40CGeBwA2EWKawrxB/49QJFRMpzSGmRHaP4aSVRZR",
    "Aj+iTrQIJzZifU7u7QmnkSc4zaJxcgg7tFlkGlmkMz4O921kW0hgrqz/6EwZcUwk1okhrvzEQjfFpTpP",
    "F0FAbQm/sZRHQnMLda2Qa6GwFXJH/QedkGuhtBVys8iwMXxQxduyttTeltGleFu2zxLHmu2leFu2l+Jt",
    "2V6Kr7d9jxXvnZ8GFuqY84Dxnyj6Z4Oxfx8RqO/qdS4ZRqmixbx1Nq8rZ+KdWxKT+5jMZw6f/u9f1afM",
    "e3DurAy+CtHWRo+Cvs7BLn60qLd1587eycVbo4ACwBYJFAC2aKAAqCeCTj64hmEYr9Wf/yBD6/Fiktrh",
    "mjEJFtlMRP85DD9HWoj+Vfo+82OoT21NAdXj2Ei3Gz7RxAPJCumvnsNCPbwSboElVsIzH9uwfQXDxnME",
    "8BOxpc7s+cOcxjAl8br/9H3GgoAtqWeGKFUNXX7cHqcx05Q9PaJchvMZSfykf4NdyGUxzjWZ9y/9NoAl",
    "H5ai6HIL1pMEjrk+7NRfyvn6n3xOJz/t3zbP766vnDOYtokeQlvSbU29Ct3PfRtdZSaaeRZ6YSEaBgN+",
    "BLNwzMKEsQD4OX2YMALLpeqm1EucoIxXsrU5miU0QvAtzIGKhUAprRfeKdgFxJiEcxujQyH8Dmh/CROg",
    "NuagBcCvCSzjgt81ai2vrlbKbK2uTnLu9A27dA/K9H6ymLyirmbYqQSCsnJKKAUMx52i+yF9g5aaGmqD",
    "lprCwtTyPCCwWFH7u/8mTRH65qgdFNaM6IwKs4DF00WA8c153hbztHlbzOOyYBFGCVJl0RSnsWjaQWGc",
    "fwSqZixtcu3PYt/D2Ei0QxhItENYR7RDmEa0w9pFs0Ck0Z6aFSCN7TRLPEztxHRW/VpQtSeokK1oh/Cf",
    "aIfwn2iH8J9oh/CfaIfwn2iH8J9oh/Df3oVDp1PoL1FMpbRG+FJpjfAon8Gj4RyWpccPjZVIJf4uA3pP",
    "dJPTpoC/jdmUvybAIs1aa2PY88k6ZI2RtUSYGYZVGIblzZBoGm92KfCeEphwgBXzdqb9swKQc3X7SBKv",
    "LNQ3a3rg7M0DzVjkyr+fpc541mHq70D8OGOUz3kDqfwuf8nBKFxv0EbL7BmEX1PPX4S5aXSx3Qgx2hxC",
    "kweNEPvNEIK4NR1oo3x4pUvrAmklIR+r/2GzfFHgYPWH996a9BfysfqL93A0ISrtI+QjCetQdCsa+Rfw",
    "Gp6DT69DU+4Wg5pO9HBoyuACosMjmJK4kN+BJEzmX6NPmP9zYS0wmipMvsiARJp1RDG5I0MRwdoRxeSU",
    "MrN2tVsbiu2KtTHXmoFKM4dNFGycThRBUayGQHazG3O8+bkaJy0PNib7rkAbs35XoI3pvyPQZv1AVxAT",
    "CRVsKjuErlgmKiqweuC8QxMbFUB90F7rvgJbc5jcVO0rsCgmB1X7CiyKyTu6vgKLhekrsFit+wosUGvy",
    "xgK1Jm8sUGvyxgK1Jm8kUDvyxoKYWKHguRJ5Y7FM3FBgqeSNBTLRQwGkkjcSaPN5krzfQw7oNpww6Ypi",
    "clCVvLHPYvKOjryxWBjyxmK1Jm8sUGvyxgK1Jm8sUGvyxgK1Jm8kUDvyxoJgyBuLZeKGglNV8sYCmeih",
    "AFLJGwnUnrw1P9M1DcVbkjcWxeSgKnljUUze0ZE3FgtD3lis1uSNBWpN3lig1uSNBWpN3lig1uSNBGpH",
    "3lgQDHljsUzcUHCqSt5YIBM9FEAqeSOB2pO3Zq1Ez+SNRTE5qEreWBSTd3TkjcXCkDcWqzV5Y4FakzcW",
    "qDV5Y4FakzcWqDV5I4HakTcWBEPeWCwTNxScqpI3FshEDwWQSt5IoPbkrVmw1jN5Y1FMDqqSNxbF5B0d",
    "eWOxMOSNxWpN3lig1uSNBWpN3lig1uSNBWpN3kigduSNBcGQNxbLxA0Fp6rkjQUy0UMBpJK3AIIti9Xd",
    "h/kWvWJnb1j+lMLrOqeDeb5nAl8RBXsR832V5ebC4sYXYvth3o6/TgP3vCGwPbS65a9cuCDehlqtQc3v",
    "3MlWH8AOyVzGaz9K2OtF6SLsnMTGc+LSi0vtlZvyFe8VbC/9K75WEV7ELV1MIjK/Y8IYUmGpBt8C6iyA",
    "HSf4TuF5owlJKN9Ai2sK9pKqwr7SXOO47U7S+X7RIH1t82gwvE4rl7/ukWuzI/5lVkve5Wd35Uq35N05",
    "39ZaOCE/FxDYFlmeo9HWyzF3Zb6x9eng3Wzr/IafmvgebH1N4q3xmXxU8YzwyMLT1dhwZxAcLn+fC1pr",
    "YuNAbC6txoZpN47VkrUsXNY6eHHKHE3CCnqFxV5cBmX3K8rKV0CVt8OaVawGOSzGEcqnkyCLGvjPOQ2C",
    "ayJiKGVzMCBsCC/WAmR5570lWZuATnkkwtXhjuD70vUJS2EXdX37WCzX1QoAY6nKZIdcSb0V1yihIIHd",
    "iu3UvRV5Dnm/4BuLi9CU0crfAS9OyXNmB6/TBfhT7PEOf3MxfDldFpBzBnuzj4b7cnpYuUeYhIeBuOVo",
    "D75AkKV2Jg+efD3HV2k2lG8yqGmWnYNGDdlSb7a9itn4fmv5K5IbmE1a5H9vtuXJK7fkhcyqP5wxsw8A",
    "qGSjeXF2A7vKn4TXl6dtQEF5rsunrglRntFF9B0PZYevxOfaDcOjPamK7o7dw5EMS90dewcHcp5Nd8do",
    "/yjvAle5tKbH/ui4QdOD0bBB08O93QZNj3ZHDZqCwRo0He7sHDaoOtw5Pm7QdTg8Br7NaFtjkuEuqNtw",
    "y97hqEnd0cH+OgEVdZBAV6qf0vFNdrxW6YhTq0pCHNZWNVyspL31BC1XNB+++f1//vUnHrDy4xirE6si",
    "YnWOFzXyqE6XtVom1wEqjFUlIy2qUmxm5WaKheiGt79BU1d8WmTMv6hRLlYPK4R7J/bCzXp62Qk1d/Ib",
    "8W23/r6uGmjR21ebgwE79fUGqx5VrPr9395//Mu3H77944c/fPHpy/cfv/jtx6++zpKlnn03sqim6obP",
    "zNA425v6mkVMCdaaK6uorbnIw7d8OovjVYTKhF6LUHGuOULri4DjivXEiPC2GHsJDTBmK/VGS99jS/6C",
    "YsyCvK+WTyOHXgKpO+X4Ysw15ZtdXcHQCQqxXYmzKhIglmeSPeuGXTk5gFFNBKV+mec5/3pP4fzaK7nz",
    "1YuXax/0SZVP/TyFIVGmYolSNTT2msZFjSv7l7UoEeeao2RtYGXIumH26R614Pn+r//89NXvnA//+POn",
    "r/9em26ya1Rrb1kiVwcu0mflrMuNmI9lS0Rfiqw1W9UPX7vU1Rsba1jJMjDWd++/bGks6df/c2NVh3Oi",
    "lAZW/+6bf2fE3tJusl+v2i3nqHKUyW9zrUoLXP2x4u0WlQXkqBjZJk/+C1BLAwQUAAAACACHTuJArU61",
    "37IOAAAwPgAAEQAAAHdvcmQvc2V0dGluZ3MueG1stVtbbxvJlX5fIP/B0LutuleXEE1Q12QGlie7shOM",
    "3yiqZXHNG5qUFfvX79ekaHnGXw8GG+SJzT5dt3Or75xT9ee//Gu1fPGpH3aLzfryTL4SZy/69Xxzu1h",
    "uDx797a97M5e7Paz9e1suVn3l2ef+93ZX37403/9+fFi1+/3+Gz3Al2sdxer+eXZ/X6/vTg/383v+9Vs",
    "92qz7dcg3m2G1WyPv8OH89Vs+PiwfTnfrLaz/eJmsVzsP58rIdzZUzeby7OHYX3x1MXL1WI+bHabu/3Y",
    "5GJzd7eY908/pxbDHxn32LJs5g+rfr0/jHg+9EvMYbPe3S+2u1Nvq/9vb1ji/amTT7+3iE+r5em7Ryl+",
    "78un5T5uhtuvLf7I9MYG22Ez73c7CGi1PC53NVusv3YjzXcdfWX1K7D6/Dj2+dgVmktxeHqe+W75XXsi",
    "7aMUXy9uhtlwFDMU4JtZbHf5YbffrMpsP/va3+Pj46vH7e7VfP00iW+kJvU5SM+Nzl6s5hc/flhvhtnN",
    "Eur5KM3ZD9DNL5vN6sXjxbYf5hA3FFuIs/ORcLvYbZezz2k2//hh2Dysb6/vZ9sen36aYUXy+NENFg2T",
    "KJs3m/31w3D47m/9DO/+wIdts9l/9+Htk+L9fQBxPiodeurXsIx5Pyrk5dlpfv3d7GG5fzu7ud5vtqfh",
    "jHoiz+9nw2yO/q+3szlkmzfr/bBZnr67HWecYVsDRH9cy9HSxqXv0KRvm+Hd6wMnHpZvh9liOXbUH1/s",
    "+lZfzz5vHvaH/8eW10crxwjr2QocPr59styrzW1/BtLDsPgqvpPxT6rT2ODIbfvtFH870AZeaVjc9mDF",
    "sr/ef15i7uv99eJLH9e3P0FtFvAFB/v9N2bwexPo1+PIP8OHvf287Vs/2z+Arf+hwQ6Sa8vF9moBfRt+",
    "XN9CK/5jgy3u7voBAyxm+/4K6rYYNo8HPh+1/N8d9/zx4lntoPr/mA27g/EdHr9VpdUGO8t+9lUl+s8/",
    "bd7/dfnxx//dLN6//cf9L19+FFfqzccr9Yt+X94s3peP5ue37/71/m00V1/e2Tdf5urnf/6yuPvvy8tR",
    "lhj5m/Gwq90eBh4f/gd2edI8IbrOhfRkVCP1mSKk0urJEfyWYn1XjirzG4qSzRpK0dJ0lVOsVlOUpDxv",
    "45UNlGJs7jpKsSbFRileR5c4xeaJGXgfHOdb0J3gsw66Sj5OsKry9QTXFG8Thax8nKhy45KLxlrO6+RK",
    "5ZTsu8B7K7qLkfKtaq0nKCZnzdv40KhWSaGzoZKT0sjgWG9QXimpHoBi20QbKxydtYQlRKrXUmtfKN+k",
    "tpHrmzTCaN6bEV3lPDCuZE6xTmjOHeuj4utxqljOg04kbSlHg9aOrzRal/l6oo2nLfvXnkImUQRfTzLB",
    "TFBcm+BONi5SW5BVRMm5U7VLfD3VuUptWzahZKbcadpOcLRhHK6JzQlJrV42LyVfT/NdR2WqhHGNjqOE",
    "bdyylHBKUB6AoiOdm5K2citR0mXuXZT2zlOrByVmPo5RNVEeKCtN5Su1esJOlbUmKCY5Zb1tVNrKidqo",
    "71Ve28yl0Fkd+UqDUVymKhhj+dwifAXnTtTJUU1U0Xgn6UqziZXaqSqi8R1QFVMCn0GxplHLUsVFvv+o",
    "ir2E99ZkpyYoOiYuhWYURw6q2WKpD1HN10J708LKPEFx3lMr0cJbQ/VAS68U5Y5WzoUpSsd3GQ3nwrUK",
    "248IVKbaAN05pgfayMR3DG1MMZa2sdB4anPa6s5zvjljuI/XziXN23TGFr6eznmu8Rg/J2oLOoDZXHLB",
    "AdfQlQZfDPU7OgpjqI6CUp2ivSVpuYfVCRpHPYVORhWqvaAEx+WTfVWcB8WZCR5UJRLXkAofS32Ixiia",
    "ekvdnDZ0BkZYXyYowPF0pUbKwK3eSF85QoF/HRNoh6D912jDKCcz1QOjvOWez2Ar4d7fANxy5G00dnTq",
    "D0AJHKkaCFtSmWJTCJVaibEOmyNdqXON40TjBcIC2sarzNGT8cZZqgfGe+35rDsRJ9p0WvD4x3Q6cA9r",
    "OrgQztHO2UY9BTb6WrgeBJ+5fzNRCL6bgRI4ujXRGG4LJmvvqZWY4mPkmlil4YjLNKBlLrlmJY8+TLOt",
    "o9yxcGKCyhSUXKjns0p4TS3YQqTcIwGKOa7XFn6C2481vvF9wWIDlNTDIvyxhkrbYs5VM4233ujI1wPt",
    "rby3TqlMNd52sEY+tw54h48TRORowwYZeN7FBj/hXUCRgVqJRZaA5w8AKDSPPmz0ZoLXySq+z9lkXeRa",
    "lVzjFmyTz5q3yd5zW7BVIl1EZVpN4V7MVgQMnDugeD6DBj3g4zTpRKIzaIiNJtogW0Qt2CEC4+jJSSkS",
    "3Rud1NlRO3UwuUQRisMOGKgtOA0T5nPrRAvUi7lOBu53XBBKUFwFSmvUFlxQwDWMoy4YwVGnC04HPk5U",
    "3cTcoG88N+iy6hqVHFAigBqdW/ZNcvkgmkq8t6JqnGijFcewrsBbcvkUJDCoJjpkYbnVe2EUz355YR2P",
    "Dr3w3lH5eKkm/LVHYozvzqA0jsW80oVjCq/MhHyAQjTXUW8tgj0mOe+E4dlW77Ch85U6n7i+eQ+YRqUw",
    "co3nqzyQUKMeCZSJLBsojiMHUPyE5ICeeLztA4JDig98kJbbj4edWmpzPiLPxnkALMi110d4PmolPplk",
    "+NySrxPSztpyZOcLfCK1H1985pl63+B2+EqbmMhBwhvoxGWKHYMj7064iaxHhyy1ojPoEDEVqqMd+Mnr",
    "MqDYQr0YKI7n7DptygTF6tbRfaGzpiiKkTprC0dPnUVqkMoHAX/ifqfzorN0L+m8yTzbOl3p6pC/FnzW",
    "SSieLeqSMDzb2iU9kWXrkgWNeSRQMs8WddkKvjN1GXs6l0LRbaINsjsc4XfFFU3jkq4CilGM1MFKFLXg",
    "IEQOtLeAGIP7+ICCCY+zAiJ+nnUPCmUexzgakNbNnDLG9dQWgraJa2/QTkveBlnqiVkjxOB5WKR7paV6",
    "EKy0PEIGpeOYApTI82LBIrlOvXJArcJQ3Qko7XLsH1AqiFzasNOJuQFB8hgjIBLveG/Iv/H4NEShC581",
    "MnOa+sQQgckn2iBrR31IiC5yfB2SFLyKAAocD9XEpAuPaUMyYDZv4xPPGoYiM6/lhKICz2CEYmJHPWxA",
    "dYxnpUKVlVfyQ1WKo/VQdeXYP1QjJuRTXeVoI1TveCwDSuLZiNBc4fFclELxSj4UZIJvUemOY/KonG9U",
    "30DpeBSKcDtwbBmR3ec5b1Dy6cTWrzOa0SAFST0sKKh+MK1CusryHHFE4ZvXCrA1IvNCe0PomihKg9sT",
    "XK+j04l7l+hl5OdQxmQEz7JF1Nr4fhoDKncU+2Mrc5WvJ/jAK5ExjvkVygNk9zkKiEkjBcfbWMsxecwi",
    "8axHxD7LLRjJKpn53JqyvAoXm0Gukc0tCTnh/ZPQjs8gIbsfqL6BEjmmSAI5DOqRkvQd5xuOAkUefSSE",
    "TJHyOmmnwgQF2JLia4Qeme9mydiOV6STcQCKlKPYt/mOjlSamRgHgStH68lBD+hulhwqJnQ3Sx75L75S",
    "FOU76sVSJ/KEFDoca+HjIB/PI+QUUOWn6CkFPVGbSsHYyGcdrOAZwBSwl1BMkVCf47aQonQTehChitTD",
    "pmgNjw4TzGSCb/B7/ExJyq5xvJMK6vKc10VHHjunghwxnzWiUE8RcWpuosKeAbgctR9QGl9PRoaJzy3j",
    "dFWa6A2FJrrSDEDMMXlG5pZHu1lJJai+4TBfU1RDUNRUPEudEQfzemNGVMA5iqAtG6q92UrNUXS2TvOM",
    "DAI9HBpg3gWUlqk/QALd85NXudOS5w9QBERRgo6DaJdXRXJGBoPutLmMTKC9FRysoVg5F6Qg6QwQEnT8",
    "PF/BGUCODwrO63O/gyTFBIou4+kMPgMjBT+XhuMCmde3RxbwSmSxOLhIbQHHaoB8Gd+Kg5JS7IISvyy8",
    "N68VRxsFOXyeZSsd/A6fAbKGilpWCWbiNCoobmKl8Nf8zFxJMvCTIwWZer4DlmwSj9GReAocV5UMH9Io",
    "r5Fx4Pi6VOxZVHtLxUlMisVKE3VCCg3nD+gMKs4GclRTJSop1OqrxNyofCpyRTzvUrVEh4wHFSUJfoIV",
    "4ZziCAUUHAOhvY0U6hMrfD+3n2o8sB3tDYfPOIKESslCpQCo7Hk2r3qJ4xl0HOTj+f5TgzKW6kENNvCI",
    "vwbUujivA6Jd6kcrgCKPcmrC/kOtHpSJ82KgZF5FqFmhaER5kHGohMs0oz7GKUUq7i1r8XWiTRWe75q1",
    "msTr2xVC5VW42gA6OXcagDe3koYgh8q0jYfpqO9tAhl0qm9NCsu9WENNnttPU2Li3BOKy557/4ajfvzs",
    "JA5I5laZTJvBWTJqWThSovmO3pyyvJLSvFMcdbZOCW4/rbOJ47eGOym8jt6CnDiD0VD17ahMW0LKjtpc",
    "y8hoUqtvAHDcI7VRr7keFO8K9W+tapEzlQLOhlNP3inwjWZOI/Y5nOkgvSWZBWAfodQmkGkkei1VLKay",
    "U/AStUsE1URDFGoLVkRCwblSNBPEhyBZlUxj4+BYtsa1V8JRh6IvSuLl+/V41OQ9ArqRgotOY2IMt5tW",
    "F+PN078Pp6fxltyL1fGGXZ6tbobF7MXVeDcVrVYXN8PHtFif6Dc9biD231KuH25OxJcvj4TdarZcNlw7",
    "PBEOSrU6XKUs/d2h2+XVbPjw3O/TFwN9e9vf/fS1r/F6Zj/8FVcst8fRHofZ9nj77TScxLmtI22x3r9e",
    "rE7vdw8316dWa9wv/YaEe50/fxrGRufP7Hm82ONa8uEW4evZ+sPpxle/fvnueryA1s92+7hb4CLql/uX",
    "+c3YGvfYlsP1eJu5v5ptt7hwie9uPsjLs+Xiw/1ejs32+HeLW82HPzcf1BNNHWj4N9IOf2bzcbH4+ulh",
    "/OD4iK+eHp7f6dM7/fwOd3aP35nnd/b0zj6/c6d3uFX9eHGP24vDcrH+iCuap8fx/d1mudw89rd/O728",
    "PPvu1YGFz1e9f/g/UEsDBAoAAAAAAIdO4kAAAAAAAAAAAAAAAAALAAAAd29yZC90aGVtZS9QSwMEFAAA",
    "AAgAh07iQIDADf+SBgAAiBsAABUAAAB3b3JkL3RoZW1lL3RoZW1lMS54bWztWU9vG0UUvyPxHUZ7b2Mn",
    "dhpHdarYsRto00axW9TjeHe8O/XszmpmnNQ31B6RkBAFcaASNw4IqNRKXMqnCRRBkfoVeDOzu96J1ySB",
    "CCpoDq09+5v3//3mzfrqtQcxQ4dESMqTtle/XPMQSXwe0CRse3eG/UsbHpIKJwFmPCFtb0akd23r3Xeu",
    "4k0VkZgg2J/ITdz2IqXSzZUV6cMylpd5ShJ4NuYixgq+inAlEPgI5MZsZbVWW1+JMU08lOAYxN4ej6lP",
    "vK1cbI+B7ERJveAzMdBCySI2mNQ1Qs5klwl0iFnbAw0BPxqSB8pDDEsFD9pezfx5K1tXV/BmtompJXtL",
    "+/rmL9uXbQgmq0anCEeF0nq/0bqyU8g3AKYWcb1er9urF/IMAPs+eGptKcts9DfqnVxmCWQ/Lsru1pq1",
    "hosvyV9bsLnV6XSarcwWK9SA7MfGAn6jtt7YXnXwBmTxzQV8o7Pd7a47eAOy+PUFfP9Ka73h4g0oYjSZ",
    "LKB1Qvv9THoBGXO2WwnfAPhGLYPPUVANRXVpFWOeqGW1FuP7XPQBoIEMK5ogNUvJGPtQv10cjwTFWgHe",
    "JLj0xC75cmFJ60LSFzRVbe/9FEMvzOW9fvHt6xfP0PHD58cPfzh+9Oj44fdWkLNrFydhederrz/5/cmH",
    "6LdnX716/Fk1XpbxP3/30U8/floNhPaZm/Py86e/PH/68ouPf/3mcQV8W+BRGT6kMZHoFjlCBzwGx0xU",
    "XMvJSJxvxzDCtLxjOwklTrDWUiG/pyIHfWuGWZYdx44OcSN4VwB9VAGvT+87Bg8iMVW0QvONKHaAe5yz"
]

# --- 移動端 App 化支持 (PWA) ---
pwa_html = """
<link rel="manifest" href="https://raw.githubusercontent.com/manus-agent/pwa-manifest/main/manifest.json">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
<meta name="apple-mobile-web-app-title" content="PDF工具">
<link rel="apple-touch-icon" href="https://cdn-icons-png.flaticon.com/512/4726/4726010.png">
<style>
    .stButton>button { width: 100%; border-radius: 10px; height: 3em; background-color: #007AFF; color: white; font-weight: bold; }
    .stMetric { background-color: #f0f2f6; padding: 10px; border-radius: 10px; margin-bottom: 10px; }
</style>
"""

# --- 公共函數部分 ---

def extract_values_from_filename(filename):
    values = re.findall(r'\d+', filename)
    if len(values) >= 3:
        return values[:3]
    return None

def extract_table_value(pdf_path, page_num, row_num, col_num):
    try:
        tables = camelot.read_pdf(pdf_path, pages=str(page_num), flavor='stream')
        for table in tables:
            df = table.df
            try:
                value = df.iat[int(row_num), int(col_num)].replace(',', '').replace(' ', '')
                return value
            except IndexError:
                continue
    except Exception:
        pass
    return "N/A"

def extract_row_values(pdf_path, page_num, keyword):
    try:
        tables = camelot.read_pdf(pdf_path, pages=str(page_num), flavor='stream')
        for table in tables:
            df = table.df
            for i, row in df.iterrows():
                if keyword in row.to_string():
                    values = [val.replace(',', '') for val in re.findall(r"[\d,.]+", row.to_string())]
                    return values
    except Exception:
        pass
    return []

def add_thousand_separator(value):
    try:
        value = float(value)
        if value.is_integer():
            formatted_value = "{:,.0f}".format(value)
        else:
            formatted_value = "{:,.1f}".format(value)
        return formatted_value
    except ValueError:
        return value

def evaluate_expression(expression, values):
    for key, value in values.items():
        expression = expression.replace(f"{{{key}}}", str(value))
    try:
        result = eval(expression, {"__builtins__": None}, {})
        return add_thousand_separator(result)
    except Exception:
        return "N/A"

def replace_and_evaluate_in_run(run, values):
    full_text = run.text
    for key, value in values.items():
        placeholder = f"{{{key}}}"
        full_text = full_text.replace(placeholder, str(value) if value is not None else "N/A")
    expressions = re.findall(r'\{\{[^\}]+\}\}', full_text)
    for expr in expressions:
        expr_clean = expr.strip("{}")
        result = evaluate_expression(expr_clean, values)
        full_text = full_text.replace(expr, result)
    run.text = full_text

def replace_and_evaluate_in_paragraph(paragraph, values):
    for run in paragraph.runs:
        replace_and_evaluate_in_run(run, values)

def process_word_template(template_stream, values, remove_text_start=None, remove_text_end=None, extra_removals=None):
    doc = Document(template_stream)
    for paragraph in doc.paragraphs:
        replace_and_evaluate_in_paragraph(paragraph, values)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_and_evaluate_in_paragraph(paragraph, values)
    if remove_text_start and remove_text_end:
        delete_specified_range(doc, remove_text_start, remove_text_end)
    if extra_removals:
        for start, end in extra_removals:
            delete_specified_range(doc, start, end)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def delete_specified_range(doc, start_text, end_text):
    paragraphs = list(doc.paragraphs)
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(paragraphs):
        if start_text in p.text:
            start_idx = i
        if end_text in p.text and start_idx != -1:
            end_idx = i
            break
    if start_idx != -1 and end_idx != -1:
        for i in range(end_idx, start_idx - 1, -1):
            p = paragraphs[i]._element
            p.getparent().remove(p)

def convert_docx_to_pdf(docx_bio):
    with open("temp_output.docx", "wb") as f:
        f.write(docx_bio.getbuffer())
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "temp_output.docx"], check=True)
    with open("temp_output.pdf", "rb") as f:
        pdf_data = f.read()
    return pdf_data

# --- 儲蓄險特有邏輯 ---

def find_page_by_keyword(pdf_path, keyword):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and keyword in text:
                    return i + 1
    except Exception:
        pass
    return None

def get_value_by_text_search(pdf_path, page_num, keyword):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            page = pdf.pages[page_num - 1]
            text = page.extract_text()
            if not text: return "N/A"
            lines = text.split('\n')
            for line in lines:
                if keyword in line:
                    matches = re.findall(r'[\d,]+', line)
                    nums = [m.replace(',', '').strip() for m in matches if m.replace(',', '').strip().isdigit()]
                    if nums: return nums[-1]
    except Exception:
        pass
    return "N/A"

def extract_values_from_filename_code1(filename):
    values = re.findall(r'\d+', filename)
    if len(values) >= 6:
        return values[:6]
    return None

def extract_nop_from_filename(filename):
    values = re.findall(r'\d+', filename)
    if len(values) >= 11:
        return values[5], values[7], values[10]
    return None, None, None

def extract_numeric_value_from_string(string):
    numbers = re.findall(r'\d+', string)
    return ''.join(numbers) if numbers else "N/A"

# --- Streamlit 界面 ---

st.set_page_config(page_title="PDF 計劃書工具", layout="centered")
components.html(pwa_html, height=0)

st.title("📄 PDF 計劃書工具")

menu = ["儲蓄險", "儲蓄險添加", "一人重疾險", "二人重疾險", "三人重疾險", "四人重疾險"]
choice = st.selectbox("選擇功能類型", menu)
export_format = st.radio("選擇導出格式", ["Word (.docx)", "PDF (.pdf)"], horizontal=True)

with st.expander("📁 上傳文件", expanded=True):
    template_file = None
    if "重疾險" in choice:
        template_file = st.file_uploader("上傳 Word 模板 (.docx)", type=["docx"])
    if choice in ["儲蓄險", "儲蓄險添加"]:
        pdf_file = st.file_uploader("選擇連續提取 PDF", type=["pdf"])
        new_pdf_file = st.file_uploader("選擇分階段提取 PDF (可選)", type=["pdf"])
    else:
        num_files = {"一人重疾險": 1, "二人重疾險": 2, "三人重疾險": 3, "四人重疾險": 4}[choice]
        pdf_files = []
        for idx in range(num_files):
            pdf_files.append(st.file_uploader(f"選擇第 {idx+1} 個 PDF", type=["pdf"], key=f"pdf_{idx}"))

if st.button("🚀 開始處理"):
    if "重疾險" in choice and not template_file:
        st.error("請先上傳 Word 模板！")
    else:
        with st.spinner("正在處理中..."):
            if choice in ["儲蓄險", "儲蓄險添加"]:
                if not pdf_file:
                    st.error("請上傳 PDF 文件！")
                else:
                    with open("temp_pdf.pdf", "wb") as f:
                        f.write(pdf_file.getbuffer())
                    filename_values = extract_values_from_filename_code1(pdf_file.name)
                    if not filename_values:
                        st.error("PDF 文件名格式不正確。")
                    else:
                        target_page = find_page_by_keyword("temp_pdf.pdf", "退保價值之説明摘要") or 6
                        doc_fitz = fitz.open("temp_pdf.pdf")
                        page_num_g_h = len(doc_fitz) - 6
                        g = extract_table_value("temp_pdf.pdf", page_num_g_h, 11, 5)
                        h = extract_table_value("temp_pdf.pdf", page_num_g_h, 12, 5)
                        s = extract_numeric_value_from_string(extract_table_value("temp_pdf.pdf", page_num_g_h, 11, 0))
                        i = get_value_by_text_search("temp_pdf.pdf", target_page, "@ANB 56")
                        j = get_value_by_text_search("temp_pdf.pdf", target_page, "@ANB 66")
                        k = get_value_by_text_search("temp_pdf.pdf", target_page, "@ANB 76")
                        l = get_value_by_text_search("temp_pdf.pdf", target_page, "@ANB 86")
                        m = get_value_by_text_search("temp_pdf.pdf", target_page, "@ANB 96")
                        pdf_values = {"g": g, "h": h, "i": i, "j": j, "k": k, "l": l, "m": m, "s": s}
                        values = dict(zip("abcdef", filename_values))
                        values.update(pdf_values)
                        remove_start, remove_end = None, None
                        extra_removals = []
                        if choice == "儲蓄險添加":
                            extra_removals.append(("信守明天多元货币储蓄计划概要：", "信守明天多元货币储蓄计划概要："))
                            extra_removals.append(("(保诚保险收益最高的储蓄产品，", "适合身体抱恙不能买寿险人士。"))
                        if new_pdf_file:
                            with open("temp_new_pdf.pdf", "wb") as f:
                                f.write(new_pdf_file.getbuffer())
                            n, o, p = extract_nop_from_filename(new_pdf_file.name)
                            new_doc_fitz = fitz.open("temp_new_pdf.pdf")
                            p_q_r = len(new_doc_fitz) - 6
                            q = extract_table_value("temp_new_pdf.pdf", p_q_r, 11, 5)
                            r = extract_table_value("temp_new_pdf.pdf", p_q_r, 12, 5)
                            s_new = extract_numeric_value_from_string(extract_table_value("temp_new_pdf.pdf", p_q_r, 11, 0))
                            values.update({"n": n, "o": o, "p": p, "q": q, "r": r, "s": s_new})
                        else:
                            remove_start = "在人生的重要阶段提取："
                            remove_end = "提取方式 3："
                        
                        # 重新組裝 Base64 字符串
                        full_base64 = "".join(TEMPLATE_CHUNKS)
                        template_stream = io.BytesIO(base64.b64decode(full_base64))
                        
                        try:
                            output_docx = process_word_template(template_stream, values, remove_start, remove_end, extra_removals)
                            if "PDF" in export_format:
                                pdf_data = convert_docx_to_pdf(output_docx)
                                st.success("✅ 處理完成！")
                                st.download_button("📥 下載 PDF 文件", pdf_data, file_name="output.pdf", mime="application/pdf")
                            else:
                                st.success("✅ 處理完成！")
                                st.download_button("📥 下載 Word 文件", output_docx, file_name="output.docx")
                        except Exception as e:
                            st.error(f"處理模板時出錯: {str(e)}")
                            st.info("這通常是由於內置模板數據損壞。請聯繫開發者重新生成代碼。")

            elif "重疾險" in choice:
                if not all(pdf_files):
                    st.error("請上傳所有 PDF 文件！")
                else:
                    all_values = {}
                    suffixes = ["", "1", "2", "3"]
                    for idx, pdf in enumerate(pdf_files):
                        suffix = suffixes[idx]
                        temp_name = f"temp_pdf_{idx}.pdf"
                        with open(temp_name, "wb") as f:
                            f.write(pdf.getbuffer())
                        fn_vals = extract_values_from_filename(pdf.name)
                        if fn_vals:
                            all_values.update(dict(zip([f"a{suffix}", f"b{suffix}", f"c{suffix}"], fn_vals)))
                        d_vals = extract_row_values(temp_name, 3, "CIP2") or extract_row_values(temp_name, 3, "CIM3")
                        d = d_vals[3] if len(d_vals) > 3 else "N/A"
                        tables_p4 = camelot.read_pdf(temp_name, pages='4', flavor='stream')
                        num_rows_p4 = tables_p4[0].df.shape[0] if tables_p4 else 0
                        e = extract_table_value(temp_name, 4, num_rows_p4 - 8, 8)
                        f = extract_table_value(temp_name, 4, num_rows_p4 - 6, 8)
                        g = extract_table_value(temp_name, 4, num_rows_p4 - 4, 8)
                        h = extract_table_value(temp_name, 4, num_rows_p4 - 2, 8)
                        all_values.update({f"d{suffix}": d, f"e{suffix}": e, f"f{suffix}": f, f"g{suffix}": g, f"h{suffix}": h})
                    output_docx = process_word_template(template_file, all_values)
                    if "PDF" in export_format:
                        pdf_data = convert_docx_to_pdf(output_docx)
                        st.success("✅ 處理完成！")
                        st.download_button("📥 下載 PDF 文件", pdf_data, file_name="output.pdf", mime="application/pdf")
                    else:
                        st.success("✅ 處理完成！")
                        st.download_button("📥 下載 Word 文件", output_docx, file_name="output.docx")

st.markdown("---")
st.caption("💡 提示：儲蓄險功能已內置模板，直接上傳 PDF 即可。")
